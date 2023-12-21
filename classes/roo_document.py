import sys
import os
import copy
import json
import shutil
import csv
from dotenv import load_dotenv
from docx.api import Document

from classes.rule_set_modern import RuleSetModern
from classes.rule_set_legacy import RuleSetLegacy
from classes.rule_set_chapter import RuleSetChapter
from classes.comm_code_validator import CommCodeValidator
from classes.environment_variable import EnvironmentVariable
from classes.error import Error
from classes.warning import Warning
import classes.globals as g
import classes.functions as func


class RooDocument(object):
    def __init__(self, psr_source_file=None):
        g.rule_ends_with_or = []
        self.psr_source_file = psr_source_file
        self.get_config()
        self.get_footnotes()
        self.get_all_rules_with_classes()
        self.get_commodities()
        self.open_psr_source_document()
        self.get_document_type()
        self.validate_psr_table()
        self.read_psr_table()
        self.process_psr_table()
        if self.modern:
            self.process_subdivisions()
            self.remove_invalid_entries()
        self.sort_by_minmax()
        self.write_table()
        self.kill_document()

        self.check_opening_dash_in_rule()
        self.check_multiple_manufacture()
        if self.check_coverage:
            self.check_commodity_coverage()
        if self.validate_min_max:
            self.validate_min_max_values()
        self.write_report_on_rules_ending_with_or()

    def write_report_on_rules_ending_with_or(self):
        folder = os.path.join(os.getcwd(), "resources", "temp", "agreements")
        filename = self.docx_filename.replace(".docx", ".json")
        os.makedirs(folder, exist_ok=True)
        filename = os.path.join(folder, filename)
        if os.path.exists(filename):
            os.remove(filename)

        if len(g.rule_ends_with_or) > 0:
            with open(filename, 'w') as f:
                json.dump(g.rule_ends_with_or, f, indent=4)

        print("\nFinished processing {file}\n".format(file=self.docx_filename))

    def check_multiple_manufacture(self):
        if len(g.multiple_manufacture) > 0:
            Warning("The word 'Manufacture' appears more than once in these rules. This may be okay, but check, in case these rules should have been split.\n\n{multiple_manufacture}".format(
                multiple_manufacture=", ".join(g.multiple_manufacture)
            ))

    def check_opening_dash_in_rule(self):
        if not self.modern:
            errors = []
            for rule_set in self.rule_sets:
                for rule in rule_set.rules:
                    if rule["rule"][0:1] == "-":
                        errors.append(rule_set.heading)
            if len(errors) > 0:
                Error(
                    "Rule starts with a hyphen ... Check for incorrect cell merging, or for lines that start erroneously with 'or' on headings {headings}".format(headings=", ".join(errors)), show_additional_information=False
                )

    def get_config(self):
        """ Works out paths and other configuration settings """

        load_dotenv('.env')
        self.resources_folder = os.path.join(os.getcwd(), "resources")
        self.source_folder = os.path.join(self.resources_folder, "source")
        self.export_folder = os.path.join(self.resources_folder, "export")
        self.config_folder = os.path.join(self.resources_folder, "config")
        self.defaults_folder = os.path.join(self.resources_folder, "defaults")

        # Get paths from environment variables
        self.preferred_code_list_file = EnvironmentVariable('preferred_code_list_file', 'string', permit_omission=True).value
        self.ott_prototype_path = EnvironmentVariable('ott_prototype_path', 'string', permit_omission=True).value
        self.all_rules_path = EnvironmentVariable('all_rules_path', 'string', permit_omission=True).value

        # Get features
        self.validate_psr_tables = EnvironmentVariable('validate_psr_tables', 'int', permit_omission=False).value
        self.check_coverage = EnvironmentVariable('check_coverage', 'int', permit_omission=False).value
        self.validate_min_max = EnvironmentVariable('validate_min_max', 'int', permit_omission=False).value
        modern_documents = EnvironmentVariable('modern_documents', 'string', permit_omission=False).value
        self.modern_documents = modern_documents.split(",")

        if self.psr_source_file is not None:
            self.docx_filename = self.psr_source_file
            self.docx_filepath = os.path.join(self.source_folder, self.docx_filename)
        else:
            if len(sys.argv) > 1:
                self.docx_filename = sys.argv[1]
                self.docx_filepath = os.path.join(self.source_folder, self.docx_filename)
            else:
                Error("Please supply an input document.", show_additional_information=False)
        g.docx_filename = self.docx_filename

        # Export paths
        self.export_filename = self.docx_filename.replace(".docx", "").replace(" ", "-").lower()
        self.export_filename = self.export_filename.replace("_psr", "")
        self.export_filename = self.export_filename.replace("-psr", "")
        self.export_filepath = os.path.join(self.export_folder, self.export_filename) + ".json"

        # Configuration / footnotes
        self.footnotes_filename = os.path.join(self.config_folder, "config-" + self.export_filename + ".json")

    def get_footnotes(self):
        """ Gets a list of the footnotes associated with the selected scheme
        Footnotes are rare ... they only exist on later documents, as the functionality
        did not initially exist."""
        # Load footnotes
        self.footnotes = {}
        if os.path.exists(self.footnotes_filename):
            f = open(self.footnotes_filename)
            self.footnotes = json.load(f)

    def get_all_rules_with_classes(self):
        """ Function to pull in a list of all of the rules from the XI tariff and the
        classes of rule that are associated with each of these. These are used (for example)
        in determining if a rule is wholly-obtained, and to populate the 'class' node
        in the resultant JSON file."""
        if not os.path.exists(self.all_rules_path):
            self.all_rules_path = os.path.join(self.defaults_folder, "all_rules.json")
        f = open(self.all_rules_path)
        g.all_rules_with_classes = json.load(f)
        f.close()

    def get_commodities(self):
        """ Function to retrieve all of the commodity codes (latest version) from an external CSV file
        """
        g.all_headings = {}
        g.all_subheadings = {}
        g.all_codes = []
        found_headings = []
        self.check_preferred_code_list_file_exists()
        with open(self.preferred_code_list_file) as csv_file:
            csv_reader = csv.DictReader(csv_file)
            for row in csv_reader:
                if row["Class"] == "commodity":
                    if row["Commodity code"][0:4] not in found_headings:
                        found_headings.append(row["Commodity code"][0:4])
                        g.all_codes.append(row["Commodity code"])

                if row["Commodity code"][-6:] == "000000" and row["Commodity code"][-8:] != "00000000":
                    heading = row["Commodity code"][0:4]
                    g.all_headings[heading] = row["Description"]

                if row["Commodity code"][-4:] == "0000" and row["Commodity code"][-6:] != "000000":
                    subheading = row["Commodity code"][0:6]
                    g.all_subheadings[subheading] = row["Description"]

        csv_file.close()

    def check_preferred_code_list_file_exists(self):
        """ The commodity code list is used to work out the chapters, headings and subheadings
        that are covered by the rules of origin (PSRs). If the file that is specified in the
        environment variables does not exist, then the default file, which exists within this
        repository.
        """
        # Get default commodity code list file
        self.default_code_list_file = os.path.join(self.defaults_folder, "uk_commodities_2023-10-23.csv")
        if not os.path.exists(self.preferred_code_list_file):
            self.preferred_code_list_file = self.default_code_list_file

    def export_min_max(self):
        """ This is a function that is used for debug purposes. It extracts the full list
        of headings, min and max values, so that any issues can be debugged. The file is
        put into the temp folder (in resources)."""
        filename = os.path.join(self.resources_folder, "temp", "temp.csv")
        f = open(filename, "w")
        for rule_set in self.rule_sets:
            if rule_set.min is None or rule_set.max is None:
                print("Error with min or max of None on heading", rule_set.heading)
                sys.exit()
            f.write("'" + rule_set.heading + "',")
            f.write(rule_set.min + ",")
            f.write(rule_set.max + "\n")
        f.close()

    def sort_by_minmax(self):
        """ The final extract file needs to be extracted in sequence, according to the min,
        then the max values"""
        error_count = 0
        self.export_min_max()
        for rule_set in self.rule_sets:
            if rule_set.min is None or rule_set.max is None:
                # If there is a record with a min or max of 'None', this means that the document
                # has not been formated properly and needs to be resolved.
                error_count += 1
                print("Max or min of 'None' found", rule_set.min, rule_set.max, rule_set.heading)

        if error_count > 0:
            Error("Aborting until issues are resolved in the source data file.", show_additional_information=False)

        # If we have not aborted, then sort the rules
        self.rule_sets = sorted(self.rule_sets, key=self.sort_by_max)
        self.rule_sets = sorted(self.rule_sets, key=self.sort_by_min)

    def sort_by_max(self, list):
        """ Sorts the list of rules according to the max value """
        return list.max

    def sort_by_min(self, list):
        """ Sorts the list of rules according to the min value """
        return list.min

    def remove_invalid_entries(self):
        """ If any entries have been created that are not valid, then this function removes them """
        for i in range(len(self.rule_sets) - 1, -1, -1):
            rule_set = self.rule_sets[i]
            if rule_set.valid is False:
                self.rule_sets.pop(i)

    def open_psr_source_document(self):
        print("\nBeginning processing {file}\n".format(file=self.docx_filename))
        self.document = Document(self.docx_filepath)

    def get_document_type(self):
        """ Work out whether the document is modern or legacy
        The two document types are processed and structured very differently """
        filename_compare = self.docx_filename.replace(" PSR.docx", "")
        if filename_compare in self.modern_documents:
            self.modern = True
        else:
            self.modern = False

    def read_psr_table(self):
        print("- Reading table for file {file}".format(file=self.docx_filename))
        # Count the tables - if there is more than one then error out
        table_count = len(self.document.tables)
        if table_count > 1:
            Error("Ensure there is only one table in the document and re-run", show_additional_information=False)
        table = self.document.tables[0]

        if self.modern:
            rename_keys = {
                "Classification": "original_heading",
                "PSR": "original_rule"
            }
            required_keys = ["original_heading", "original_rule"]
        else:
            rename_keys = {
                "Heading": "original_heading",
                "Classification": "original_heading",
                "Description": "description",
                "Description of goods": "description",
                "Conditions": "original_rule",
                "PSR": "original_rule",
                "PSR2": "original_rule2"
            }
            required_keys = ["original_heading", "description", "original_rule", "original_rule2"]

        raw_table_row_data = []

        keys = None
        for i, row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)

            # Establish the mapping based on the first row
            # headers; these will become the keys of our dictionary
            if i == 0:
                keys = tuple(text)
                continue

            # Construct a dictionary for this row, mapping
            # keys to values for this row
            row_data = dict(zip(keys, text))
            raw_table_row_data.append(row_data)

        self.table_rows = []
        for item in raw_table_row_data:
            if item != {}:
                for old_key in rename_keys:
                    if old_key in item.keys():
                        new_key = rename_keys[old_key]
                        item[new_key] = item.pop(old_key)

                self.table_rows.append(item)

        sample_table_row = self.table_rows[0]
        table_is_valid = True
        for required_key in required_keys:
            if required_key not in sample_table_row.keys():
                table_is_valid = False
                break
        if not table_is_valid:
            Error(
                "The source table is not valid. Please check that all columns are labelled correctly.\n\nFor 2-column documents, ensure that there is a single table where the two columns are entitled\n\n- Classification\n- PSR\n\nFor 3- or 4-column documents, ensure the columns are entitled\n\n- Classification\n- Description\n- PSR\n- PSR2",
                show_additional_information=False
            )

    def process_psr_table(self):
        print("- Processing table for {file}".format(file=self.docx_filename))
        if self.modern:
            self.process_psr_table_modern()
        else:
            self.process_psr_table_legacy()

    def process_psr_table_modern(self):
        """
        Process the modern documents (e.g. Japan, EU, Turkey)
        """
        self.rule_sets = []
        for row in self.table_rows:
            if "Section" in row["original_heading"]:
                pass
            elif "SECTION" in row["original_heading"]:
                pass
            elif "Chapter" in row["original_heading"]:
                pass
            else:
                rule_set = RuleSetModern(row)
                # self.rule_sets.append(rule_set.as_dict())
                self.rule_sets.append(rule_set)

    def count_cells_in_columns(self):
        error = False
        print("- Counting table cells per row for file {file}".format(file=self.docx_filename))
        for row in self.table_rows:
            if len(row) > 4:
                error = True
                break
        if error:
            print("  - At last one row has an incorrect number of cells", str(len(row)))
            sys.exit()

    def process_psr_table_legacy(self):
        """
        Process the legacy documents
        """
        self.count_cells_in_columns()
        self.check_psr_table_validity()
        self.rule_sets = []
        row_index = 0
        for row in self.table_rows:
            rule_set = RuleSetLegacy(row, row_index, self.footnotes)
            # if rule_set.valid:
            self.rule_sets.append(rule_set)
            row_index += 1
            
        self.expand_subdivision_hierarchy()
        self.remove_invalid_entries()

        # Check for mixes of ex codes and non-ex codes in the heading column
        if len(g.mix_ex_non_ex_errors) > 0:
            Error(
                "Lines mix ex codes and non-ex codes, which is not permitted. Please split the following into multiple lines:\n\n{errors}".format(
                    errors=" | ".join(g.mix_ex_non_ex_errors)
                ),
                show_additional_information=False
            )

        # Check for non-contiguous codes separated by "and"
        if len(g.non_contiguous_and_errors) > 0:
            Error(
                "Lines contain non-contiguous codes separated by 'and', which is not permitted. Please split the following into multiple lines:\n\n{errors}".format(
                    errors=" | ".join(g.non_contiguous_and_errors)
                ),
                show_additional_information=False
            )

        # Check for non-contiguous codes separated by "and"
        if len(g.multiple_and_errors) > 0:
            Error(
                "headings contain multiple 'ands', which is not permitted. Please split the following into multiple lines:\n\n{errors}".format(
                    errors=" | ".join(g.multiple_and_errors)
                ),
                show_additional_information=False
            )

        # Check for non-contiguous ands in heading column
        if len(g.non_contiguous_and_errors) > 0:
            Error(
                "Lines contain non-contiguous values separated by 'and', which is not permitted. Please split the following into multiple lines:\n\n{errors}".format(
                    errors=" | ".join(g.non_contiguous_and_errors)
                ),
                show_additional_information=False
            )

        self.validate_existence_of_all_headings_subheadings()
        self.process_chapters()

        # Check for possible missing hyphens
        if len(g.possible_missing_hyphens) > 0:
            Warning(
                "Rules contain the word manufacture, but not as many hyphens as would have been expected. Bullets will be omitted:\n\n{warnings}".format(
                    warnings=" | ".join(g.possible_missing_hyphens)
                ))

    def expand_subdivision_hierarchy(self):
        """
        This is needed to transfer higher tiers of hierarchy subdivisions into the child tiers
        if the sub-tiers start with a "- " or a "- - "
        """
        rule_set_count = len(self.rule_sets)
        rule_set_grandparent = None
        rule_set_parent = None
        for index in range(0, rule_set_count):
            rule_set = self.rule_sets[index]
            if rule_set.subdivision_adoption_requirement == 2:
                rule_set.subdivision = "{grandparent}{divider}{parent}{divider}{child}".format(
                    grandparent=rule_set_grandparent.subdivision_original.strip(":"),
                    divider=g.hierarchy_divider,
                    parent=rule_set_parent.subdivision_original.strip(":"),
                    child=rule_set.subdivision_original.strip(":")
                )
            elif rule_set.subdivision_adoption_requirement == 1:
                rule_set.subdivision = "{parent}{divider}{child}".format(
                    divider=g.hierarchy_divider,
                    parent=rule_set_grandparent.subdivision_original.strip(":"),
                    child=rule_set.subdivision_original.strip(":")
                )
                rule_set_parent = copy.copy(rule_set)
            else:
                rule_set_grandparent = copy.copy(rule_set)
                rule_set.subdivision = "{child}".format(
                    child=rule_set.subdivision_original.strip(":")
                )
            rule_set.subdivision = rule_set.subdivision.replace("- - ", "- ")

    def check_psr_table_validity(self):
        """ Checks for:

        1. empty cells in the first column
        2. Check for ex being used more than once in a cell

        All of these issues need to be manually corrected in the source Word document
        """
        empty_rows = []
        double_ex_rows = []
        mixed_conjunctions = []
        more_than_one_comma = []
        row_index = 0
        last_valid_heading = ""
        for row in self.table_rows:
            original_heading = row["original_heading"]
            original_heading = original_heading.replace(";", ",")
            original_heading = original_heading.replace(u'\xa0', u' ')
            original_heading = original_heading.replace("ex ex", "ex")

            if original_heading == "" or original_heading is None:
                empty_rows.append((row_index, last_valid_heading))
            else:
                last_valid_heading = original_heading

            if self.is_double_ex(original_heading):
                double_ex_rows.append((row_index, last_valid_heading))

            if self.has_mixed_conjunctions(original_heading):
                mixed_conjunctions.append((row_index, last_valid_heading))

            if self.has_more_than_one_comma(original_heading):
                more_than_one_comma.append((row_index, last_valid_heading))

            row_index += 1

        # Priority 1 - report on empty column 1 cells
        if len(empty_rows) > 0:
            msg = "\nThere are errors in the table with empty cells in column 1. Please correct these issues before processing can complete. " + \
                "This is likely to need the empty cell to be merged in MS Word into the cell above.\n\n"
            for error_row in empty_rows:
                msg += "Row " + str(error_row[0]) + ", under heading " + error_row[1].strip() + "\n"
            print(msg)
            sys.exit()

        # Priority 2 - report on instances where ex is used more than once in a cell
        if len(double_ex_rows) > 0:
            msg = "\nThere are errors in the table with cells in column 1 that feature multiple ex codes. The system cannot work with cells such as these. " + \
                "Please split the rows into multiple rows, one per ex code.\n\n"
            for double_ex_row in double_ex_rows:
                msg += "Row " + str(double_ex_row[0]) + ", under heading " + double_ex_row[1].strip() + "\n"
            print(msg)
            sys.exit()

        # Priority 3 - report on mixed conjunctions in a cell, e.g. a comma, and, to and hyphen
        if len(mixed_conjunctions) > 0:
            msg = "\nThere are errors in the table with cells in column 1 that feature multiple conjunctions such as commas, ands, tos and hyphens. " + \
                "The system cannot work with cells such as these.\n\n" + \
                "Please split the rows into multiple rows, one per item.\n\n"
            for mixed_conjunction in mixed_conjunctions:
                msg += "Row " + str(mixed_conjunction[0]) + ", under heading " + mixed_conjunction[1] + "\n"
            print(msg)
            sys.exit()

        # Priority 4 - report on more than one comma
        if len(more_than_one_comma) > 0:
            msg = "\nThere are errors in the table with cells in column 1 that feature more than one comma or semi-colon. " + \
                "The system cannot work with cells such as these.\n\n" + \
                "Please split the rows into multiple rows, one per item.\n\n"
            for item in more_than_one_comma:
                msg += "Row " + str(item[0]) + ", under heading " + item[1] + "\n"
            print(msg)
            sys.exit()

    def is_double_ex(self, s):
        if s is not None:
            occurrence_count = s.count("ex")
            return True if occurrence_count > 1 else False
        else:
            return False

    def has_more_than_one_comma(self, s):
        if s is not None:
            occurrence_count = s.count(",")
            return True if occurrence_count > 1 else False
        else:
            return False

    def has_mixed_conjunctions(self, s):
        if s is not None:
            conjunction_count = 0
            if "," in s:
                conjunction_count += 1
            if "to" in s:
                conjunction_count += 1
            if "-" in s:
                conjunction_count += 1
            if "and" in s:
                conjunction_count += 1
            return True if conjunction_count > 1 else False
        else:
            return False

    def validate_existence_of_all_headings_subheadings(self):
        """
        Before trying to normalise the chapters, run a check to make sure that all of
        the included headings are present. If they are not, e.g. if the commodity 
        code hierarchy has changed since the rules were originally written, then
        you will get unexpected results, like with ornamental fish (0301-10) in 
        Singapore, where 030110 does not exist.
        """

        errors = []
        for rule_set in self.rule_sets:
            for heading in rule_set.headings:
                heading = heading.strip()
                if heading not in g.all_headings:
                    errors.append(heading)

            for subheading in rule_set.subheadings:
                subheading = subheading.strip()
                if subheading[-2:] == "00":
                    if subheading[0:4] not in g.all_headings:
                        errors.append(heading)
                else:
                    if subheading not in g.all_subheadings:
                        errors.append(subheading)

        if len(errors) > 0:
            msg = "The following non-existent headings or subheadings are explicitly referenced. Adjust the source document to resolve:\n\n{errors}".format(errors=errors)
            Error(msg, show_additional_information=False)

    def process_chapters(self):
        print("- Processing chapters")
        self.transfer_rule_sets_to_temporary_variable()
        # Break the full set of rules into individual chapters
        # and process them individually
        chapters = [x for x in range(1, 98) if x != 77]
        for chapter_index in chapters:
            g.residual_added = []
            chapter = RuleSetChapter(chapter_index, self.temporary_rule_sets)
            self.rule_sets += chapter.chapter_rule_sets
            if chapter.whole_chapter_rule_count > 1:
                if chapter.whole_chapter_rule_count != len(chapter.rule_sets):
                    # print(g.docx_filename, "in chapter", str(chapter_index), "has a chapter rule count of", str(chapter.whole_chapter_rule_count))
                    obj = {
                        "filename": g.docx_filename,
                        "chapter_index": chapter_index,
                        "whole_chapter_rule_count": chapter.whole_chapter_rule_count,
                        "rule_count": len(chapter.rule_sets)
                    }
                    g.multiple_chapter_rule_list.append(obj)

    def transfer_rule_sets_to_temporary_variable(self):
        """
        Empty the rule_sets variable, to be re-populated after the processing of each of the chapters
        """
        self.temporary_rule_sets = []
        for rule_set in self.rule_sets:
            self.temporary_rule_sets.append(rule_set)

        self.rule_sets = []

    def process_subdivisions(self):
        previous_ruleset = None
        for rule_set in self.rule_sets:
            if rule_set.subdivision != "":
                if previous_ruleset is not None:
                    rule_set.heading = previous_ruleset.heading
                    rule_set.min = previous_ruleset.min
                    rule_set.max = previous_ruleset.max

            if rule_set.subdivision == "":
                rule_set.subdivision = self.get_subdivision_from_heading(rule_set.heading)

            previous_ruleset = rule_set

    def get_subdivision_from_heading(self, s):
        s = s.strip()
        if "-" in s:
            return s
        else:
            if len(s) == 4:
                return "Heading " + s
            elif len(s) == 6:
                return "Subheading " + s
            elif len(s) == 2:
                return "Chapter " + s
            else:
                return s

    def copy_rule_sets_to_object_list(self):
        self.rule_set_object_list = []
        for rule_set in self.rule_sets:
            self.rule_set_object_list.append(rule_set.as_dict())

    def write_table(self):
        self.copy_rule_sets_to_object_list()
        out_file = open(self.export_filepath, "w")
        json.dump({"rule_sets": self.rule_set_object_list}, out_file, indent=6)
        out_file.close()
        dest = os.path.join(self.ott_prototype_path, self.export_filename + ".json")
        shutil.copy(self.export_filepath, dest)

    def kill_document(self):
        self.document = None

    def validate_psr_table(self):
        if self.validate_psr_tables:
            self.count_document_tables()
            self.count_document_table_row_cells()

    def count_document_table_row_cells(self):
        print("- Counting cells in each row for {file}".format(file=self.docx_filename))
        table = self.document.tables[0]
        cells = []
        cell_previous = "UNSPECIFIED"
        for i, row in enumerate(table.rows):
            cell1 = row.cells[0].text.strip()
            if cell1 == "":
                Error("Empty cell in first column not permitted - row after {cell_previous}.".format(cell_previous=cell_previous), show_additional_information=False)
            cell_previous = cell1

            cell_count = len(row.cells)
            if cell_count > 4:
                Error("There must not be more than 4 columns in the table.", show_additional_information=False)

            cells.append(cell_count)
        cell_set = list(set(cells))

        if len(cell_set) > 1:
            Error("Please ensure that all rows have the same number of columns and that they are of equal width.", show_additional_information=False)

    def count_document_tables(self):
        """ If there is more than a single table in the document, then the process stops """
        print("- Counting tables for {file}".format(file=self.docx_filename))
        table_count = len(self.document.tables)
        if table_count > 1:
            Error("Please ensure that there is only one table in the document.", show_additional_information=False)
        elif table_count == 0:
            Error("Please ensure that there is exactly one table in the document. There are currently no tables.", show_additional_information=False)

    def check_commodity_coverage(self):
        print("- Checking that all commodity codes are covered for {file}".format(file=self.docx_filename))
        self.comm_code_omissions = []
        f = open(self.export_filepath)
        json_obj = json.load(f)
        for comm_code in g.all_codes:
            v = CommCodeValidator(comm_code, json_obj)
            ret = v.validate()
            if ret:
                self.comm_code_omissions.append(comm_code)
                print("  - No coverage for commodity code {comm_code}".format(comm_code=comm_code))

        print("  - Finished validating {file}".format(file=self.docx_filename))

    def validate_min_max_values(self):
        print("- Checking min max for {file}".format(file=self.docx_filename))
        issues = []
        for rule_set in self.rule_sets:
            if rule_set.min is None or len(rule_set.min) != 10 or "," in rule_set.min or "and" in rule_set.min or rule_set.max is None or len(rule_set.max) != 10 or "," in rule_set.max or "and" in rule_set.max:
                issues.append(rule_set.heading)

        if len(issues) > 0:
            s = "  - There are issues with null min and max values - please correct:\n\n  - "
            s += "\n  - ".join(issues)
            print(s)
        else:
            print("  - All min max fine")
